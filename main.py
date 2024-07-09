import os
import datetime as dt
import tkinter as tk
from tkinter import filedialog, messagebox
from docx2pdf import convert
import docx


class InvoiceAutomation:

    def __init__(self):
        self.root = tk.Tk()
        self.root.title('Invoice Automation')
        self.root.geometry('500x600')

        #A dictionary for different banks that you can select on the invoice screen
        self.payment_methods = {
            'Main Bank': {
                'Recipient': 'SAMPLE COMPANY',
                'Bank': 'X BANK',
                'IBAN': 'XY12 3456 7890 1234',
                'BIC': 'ABCDEFGH'
            },
            'Second Bank': {
                'Recipient': 'SAMPLE COMPANY',
                'Bank': 'Y BANK',
                'IBAN': 'XY98 7654 3210 9876',
                'BIC': 'HGFEDCBA'
            },
            'Private Bank': {
                'Recipient': 'SAMPLE COMPANY',
                'Bank': 'Z BANK',
                'IBAN': 'XY12 4567 8901 1234',
                'BIC': 'ABCDABCD'
            }
        }

        #Basic labels we created for the user interface
        self.partner_label = tk.Label(self.root, text='Partner')
        self.partner_street_label = tk.Label(self.root, text='Partner Street')
        self.partner_zip_city_country_label = tk.Label(self.root, text='Partner ZIP Country')
        self.invoice_number_label = tk.Label(self.root, text='Invoice Number')
        self.service_description_label = tk.Label(self.root, text='Service Description')
        self.service_amount_label = tk.Label(self.root, text='Service Amount')
        self.service_single_price_label = tk.Label(self.root, text='Service Single Price')
        self.payment_method_label = tk.Label(self.root, text='Payment Method')

        #Input spaces after the labels where we will enter their values
        self.partner_entry = tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root)
        self.partner_zip_city_country_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_single_price_entry = tk.Entry(self.root)

        self.payment_method = tk.StringVar(self.root)
        self.payment_method.set('Main Bank')

        #Section created for choosing between banks
        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_method, "Main Bank", "Second Bank", "Private Bank")

        self.create_button = tk.Button(self.root, text='Create Invoice', command=self.create_invoice)

        padding_options = {'fill': 'x', 'expand': True, 'padx': 5, 'pady': 2}

        #In this section, we determine how Tkinter GUI elements (label, entry, dropdown and button)
        # will be placed in the window. We use the pack method to add and arrange each element in the window.
        # The pack method adds the elements to the window in a specific order and allows us to
        # set a specific alignment with the padding_options parameter.
        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)
        self.partner_street_label.pack(padding_options)
        self.partner_street_entry.pack(padding_options)
        self.partner_zip_city_country_label.pack(padding_options)
        self.partner_zip_city_country_entry.pack(padding_options)
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)
        self.service_amount_label.pack(padding_options)
        self.service_amount_entry.pack(padding_options)
        self.service_single_price_label.pack(padding_options)
        self.service_single_price_entry.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)
        self.create_button.pack(padding_options)

        self.root.mainloop()

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    def create_invoice(self):
        doc = docx.Document('template.docx')

        selected_payment_method = self.payment_methods[self.payment_method.get()]

        # It receives input from the user and stores it in a replacements dictionary.
        # This dictionary is used to replace placeholders with new values.
        # If the user enters an invalid quantity or price, it displays an error message and terminates the method.
        try:
            replacements = {
                "[Date]": dt.datetime.today().strftime('%Y-%m-%d'),
                "[Partner]": self.partner_entry.get(),
                "[Partner Street]": self.partner_street_entry.get(),
                "[Partner ZIP_City_Country]": self.partner_zip_city_country_entry.get(),
                "[Invoice Number]": self.invoice_number_entry.get(),
                "[Service Description]": self.service_description_entry.get(),
                "[Amount]": self.service_amount_entry.get(),
                "[Single Price]": f"${float(self.service_single_price_entry.get()): .2f}",
                "[Full Price]": f'${float(self.service_amount_entry.get()) * float(self.service_single_price_entry.get()): .2f}',
                "[Recipient]": selected_payment_method['Recipient'],
                "[Bank]": selected_payment_method['Bank'],
                "[IBAN]": selected_payment_method['IBAN'],
                "[BIC]": selected_payment_method['BIC']
            }
        except ValueError:
            messagebox.showerror('Error', 'Invalid amount or price!')
            return

        #It loops through all the paragraphs in the template document and
        # replaces the placeholders in each paragraph with the appropriate values.
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)

        #It loops through all tables, rows, and cells in the document and
        # replaces placeholders in cells and paragraphs with appropriate values.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        #Prompts the user to select where to save the file.
        # Then it saves the filled Word document as filled.docx and
        # converts it to PDF format with the docx2pdf library.
        save_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF Documents', '*.pdf')])

        doc.save('filled.docx')
        convert('filled.docx', save_path)

        messagebox.showinfo('Success', 'Invoice created and saved successfully')


if __name__ == '__main__':
    InvoiceAutomation()
